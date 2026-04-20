"""Tests for DOCX Phase 4B.2 — writer image emission from data: URIs.

The writer previously dropped every ``Image`` inline through the text
fallback path, so round-tripping a document that contained images lost
every picture. Phase 4B.2 adds ``_serialize_image`` which:

- Decodes ``data:image/<fmt>;base64,...`` URIs to bytes.
- Writes each image as ``word/media/imageN.<ext>``.
- Registers a ``RT_IMAGE`` relationship that the ``<a:blip r:embed>``
  attribute references.
- Emits the full ``<w:drawing><wp:inline><a:graphic>...<pic:pic>`` XML.

These tests pin that contract: URI decoding, package layout, OOXML
structure, dimension math, alt-text fallback, and end-to-end parse-write-
reparse via the DOCX reader.
"""

from __future__ import annotations

import base64
import io
import struct
import zipfile
import zlib
from pathlib import Path

import pytest
from kaos_content.model.blocks import Paragraph
from kaos_content.model.document import ContentDocument
from kaos_content.model.inlines import Image, Text
from kaos_content.traversal import find_by_type
from lxml import etree  # ty: ignore[unresolved-import]

from kaos_office.docx.reader import parse_docx
from kaos_office.docx.writer import write_docx, write_docx_bytes
from kaos_office.ooxml.namespace import (
    PIC,
    R_EMBED,
    WP,
    A,
    R,
    W,
    qn,
)

# --- Fixtures: real image bytes (not b"fake") ------------------------------


def _minimal_png_bytes() -> bytes:
    """Construct a 1x1 opaque red PNG from scratch, no Pillow dependency.

    The bytes pass Word's image sniffer because they are a real PNG:
    signature + IHDR + single compressed IDAT + IEND, CRC-correct.
    """
    sig = b"\x89PNG\r\n\x1a\n"

    def _chunk(typ: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(typ + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + typ + data + struct.pack(">I", crc)

    # IHDR: 1x1, 8-bit RGB, no interlace.
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    # IDAT: one raw scanline with filter byte 0 + RGB(255,0,0).
    raw = b"\x00\xff\x00\x00"
    idat = zlib.compress(raw, 9)
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


def _png_data_uri() -> str:
    return "data:image/png;base64," + base64.b64encode(_minimal_png_bytes()).decode("ascii")


def _minimal_jpeg_bytes() -> bytes:
    """Tiny valid JPEG — enough to pass the PIL/Word sniffers.

    Hand-crafted to start with the SOI + APP0 JFIF marker Word recognizes.
    Content is nonsense (the JPEG header is technically incomplete for a
    real renderer), but the format-sniffing logic downstream only checks
    the magic bytes.
    """
    # SOI (ffd8) + APP0 JFIF marker + EOI (ffd9). Real renderers will
    # reject it, but our _decode_image_src only looks at the MIME type,
    # and zip roundtrips bytes verbatim — good enough for structural tests.
    return b"\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00\xff\xd9"


def _jpeg_data_uri() -> str:
    return "data:image/jpeg;base64," + base64.b64encode(_minimal_jpeg_bytes()).decode("ascii")


def _doc_with_image(
    data_uri: str,
    *,
    alt: str | None = None,
    title: str | None = None,
    width: float | None = None,
    height: float | None = None,
) -> ContentDocument:
    return ContentDocument(
        body=(
            Paragraph(
                children=(
                    Text(value="Before "),
                    Image(src=data_uri, alt=alt, title=title, width=width, height=height),
                    Text(value=" after."),
                )
            ),
        )
    )


def _zip_contents(data: bytes) -> set[str]:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return set(zf.namelist())


def _zip_read(data: bytes, name: str) -> bytes:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return zf.read(name)


# --- Tests -----------------------------------------------------------------


class TestDecodeImageSrc:
    """``_decode_image_src`` is the one URI-resolution boundary."""

    def test_data_uri_png_decoded(self) -> None:
        from kaos_office.docx.writer import _decode_image_src

        expected = _minimal_png_bytes()
        data_uri = "data:image/png;base64," + base64.b64encode(expected).decode("ascii")
        out = _decode_image_src(data_uri)
        assert out is not None
        data, ext = out
        assert ext == "png"
        assert data == expected

    def test_data_uri_jpeg_normalized_ext(self) -> None:
        from kaos_office.docx.writer import _decode_image_src

        out = _decode_image_src(_jpeg_data_uri())
        assert out is not None
        _, ext = out
        assert ext == "jpeg"  # "image/jpeg" maps to ext "jpeg", not "jpg"

    def test_file_uri_reads_bytes(self, tmp_path: Path) -> None:
        from kaos_office.docx.writer import _decode_image_src

        p = tmp_path / "pic.png"
        p.write_bytes(_minimal_png_bytes())
        out = _decode_image_src(f"file://{p}")
        assert out is not None
        data, ext = out
        assert ext == "png"
        assert data == p.read_bytes()

    def test_unsupported_mime_returns_none(self) -> None:
        from kaos_office.docx.writer import _decode_image_src

        assert _decode_image_src("data:image/webp;base64,Zm9v") is None

    def test_http_uri_returns_none(self) -> None:
        from kaos_office.docx.writer import _decode_image_src

        assert _decode_image_src("https://example.com/pic.png") is None

    def test_docx_logical_uri_returns_none(self) -> None:
        from kaos_office.docx.writer import _decode_image_src

        # Phase 6.1 default is ``data:`` URI, but callers can pass a
        # builder returning a bare ``docx://`` logical URI. The writer
        # declines those (no side-channel byte store) and falls back to
        # alt text — same contract as ``http(s)://``.
        assert _decode_image_src("docx://word/media/image1.png") is None

    def test_empty_src_returns_none(self) -> None:
        from kaos_office.docx.writer import _decode_image_src

        assert _decode_image_src("") is None


class TestMediaPartEmission:
    """Writing an Image inline must produce a word/media/imageN.<ext> entry."""

    def test_png_creates_media_part(self) -> None:
        doc = _doc_with_image(_png_data_uri(), alt="a red pixel")
        data = write_docx_bytes(doc)
        names = _zip_contents(data)
        assert "word/media/image1.png" in names
        assert _zip_read(data, "word/media/image1.png") == _minimal_png_bytes()

    def test_jpeg_creates_media_part(self) -> None:
        doc = _doc_with_image(_jpeg_data_uri())
        data = write_docx_bytes(doc)
        names = _zip_contents(data)
        assert "word/media/image1.jpeg" in names

    def test_content_type_default_registered(self) -> None:
        doc = _doc_with_image(_png_data_uri())
        data = write_docx_bytes(doc)
        ct = _zip_read(data, "[Content_Types].xml").decode("utf-8")
        # The Default declaration for png must be present or Word
        # refuses to render the relationship target.
        assert 'Extension="png"' in ct
        assert 'ContentType="image/png"' in ct

    def test_relationship_registered(self) -> None:
        doc = _doc_with_image(_png_data_uri())
        data = write_docx_bytes(doc)
        rels = _zip_read(data, "word/_rels/document.xml.rels").decode("utf-8")
        assert "relationships/image" in rels
        assert "media/image1.png" in rels

    def test_two_images_numbered_sequentially(self) -> None:
        doc = ContentDocument(
            body=(
                Paragraph(children=(Image(src=_png_data_uri()),)),
                Paragraph(children=(Image(src=_jpeg_data_uri()),)),
            )
        )
        data = write_docx_bytes(doc)
        names = _zip_contents(data)
        assert "word/media/image1.png" in names
        assert "word/media/image2.jpeg" in names


class TestDrawingXmlStructure:
    """The emitted `<w:drawing>` conforms to ECMA-376 picture inline shape."""

    def test_blip_r_embed_matches_relationship(self) -> None:
        doc = _doc_with_image(_png_data_uri())
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml")
        root = etree.fromstring(xml)
        blips = root.findall(f".//{qn(A, 'blip')}")
        assert blips, "expected a single <a:blip> inside the drawing"
        rid = blips[0].get(R_EMBED)
        assert rid and rid.startswith("rId")

        rels_xml = _zip_read(data, "word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_xml)
        ids = {r.get("Id") for r in rels_root}
        assert rid in ids, f"blip r:embed={rid} must point at a registered relationship"

    def test_extent_matches_inner_ext(self) -> None:
        doc = _doc_with_image(_png_data_uri(), width=144.0, height=72.0)
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml")
        root = etree.fromstring(xml)
        extent = root.find(f".//{qn(WP, 'extent')}")
        inner = root.find(f".//{qn(A, 'xfrm')}/{qn(A, 'ext')}")
        assert extent is not None and inner is not None
        assert extent.get("cx") == inner.get("cx") == str(round(144.0 * 12700))
        assert extent.get("cy") == inner.get("cy") == str(round(72.0 * 12700))

    def test_docpr_carries_alt(self) -> None:
        doc = _doc_with_image(_png_data_uri(), alt="cat on windowsill")
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml")
        root = etree.fromstring(xml)
        docpr = root.find(f".//{qn(WP, 'docPr')}")
        assert docpr is not None
        assert docpr.get("descr") == "cat on windowsill"

    def test_graphic_data_uri_is_picture_namespace(self) -> None:
        doc = _doc_with_image(_png_data_uri())
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml")
        root = etree.fromstring(xml)
        gd = root.find(f".//{qn(A, 'graphicData')}")
        assert gd is not None
        # The DrawingML picture namespace is what tells Word to treat
        # graphicData children as <pic:pic>. Any other URI and the
        # drawing is rendered as an empty frame.
        assert gd.get("uri") == PIC


class TestUnsupportedSrcFallback:
    """When the src can't be decoded, writer falls back to alt text."""

    def test_http_src_produces_alt_text(self) -> None:
        doc = _doc_with_image("https://example.com/missing.png", alt="ALT TEXT MARKER")
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml").decode("utf-8")
        # Nothing drawing-shaped gets emitted, but the alt text survives
        # so the document remains meaningful.
        assert "w:drawing" not in xml
        assert "ALT TEXT MARKER" in xml
        assert "word/media" not in "".join(_zip_contents(data))

    def test_missing_alt_drops_image_silently(self) -> None:
        doc = _doc_with_image("https://example.com/missing.png")
        data = write_docx_bytes(doc)
        # No drawing, no alt, no media part — the Image inline vanishes.
        # Surrounding text is preserved by the run order.
        xml = _zip_read(data, "word/document.xml").decode("utf-8")
        assert "w:drawing" not in xml
        assert "Before" in xml and "after" in xml


class TestRoundTrip:
    """write → parse_docx recovers an Image with correct src dims."""

    def test_png_roundtrip_recovers_image(self, tmp_path: Path) -> None:
        orig = _doc_with_image(_png_data_uri(), alt="recovered", width=72.0, height=72.0)
        out = tmp_path / "rt.docx"
        write_docx(orig, out)
        reloaded = parse_docx(out)

        images = list(find_by_type(reloaded, Image))
        assert len(images) == 1, f"expected exactly one Image, got {len(images)}"
        img = images[0]
        # Phase 6.1: reader defaults to data: URIs so the writer can
        # round-trip. The payload is the actual PNG bytes base64-encoded.
        assert img.src.startswith("data:image/png;base64,"), f"unexpected src {img.src!r}"
        payload = img.src[len("data:image/png;base64,") :]
        assert base64.b64decode(payload) == _minimal_png_bytes()
        # Width/height survive (points → EMU → points).
        assert img.width is not None and abs(img.width - 72.0) < 1.0
        assert img.height is not None and abs(img.height - 72.0) < 1.0
        assert img.alt == "recovered"

    def test_full_roundtrip_preserves_image_through_second_write(self, tmp_path: Path) -> None:
        """Phase 6.1 contract: write → parse → write → parse loop
        preserves the image at each stage. Pre-6.1 the second write
        dropped the image to alt text."""
        orig = _doc_with_image(_png_data_uri(), alt="chain", width=72.0, height=72.0)
        a = tmp_path / "a.docx"
        b = tmp_path / "b.docx"
        write_docx(orig, a)
        loaded = parse_docx(a)
        write_docx(loaded, b)
        reloaded = parse_docx(b)
        images = list(find_by_type(reloaded, Image))
        assert len(images) == 1, "second round trip lost the image"
        assert images[0].alt == "chain"
        # Bytes survive the second write unchanged.
        assert images[0].src.startswith("data:image/png;base64,")
        payload = images[0].src[len("data:image/png;base64,") :]
        assert base64.b64decode(payload) == _minimal_png_bytes()

    def test_media_bytes_are_identical_on_roundtrip(self, tmp_path: Path) -> None:
        """The bytes we wrote are the bytes Word would unzip."""
        expected = _minimal_png_bytes()
        orig = _doc_with_image("data:image/png;base64," + base64.b64encode(expected).decode())
        out = tmp_path / "rt.docx"
        write_docx(orig, out)
        with zipfile.ZipFile(out) as zf:
            assert zf.read("word/media/image1.png") == expected


class TestReaderImageSrcBuilder:
    """``parse_docx(image_src_builder=...)`` lets callers pick the URI
    policy for reader-produced images. Mirror of
    ``kaos-pdf.extract_pdf.image_src_builder`` (kaos-pdf/1782546)."""

    def test_default_builder_emits_data_uris(self, tmp_path: Path) -> None:
        orig = _doc_with_image(_png_data_uri(), alt="default")
        out = tmp_path / "rt.docx"
        write_docx(orig, out)
        reloaded = parse_docx(out)  # no builder → default
        img = next(iter(find_by_type(reloaded, Image)), None)
        assert img is not None
        assert img.src.startswith("data:image/")

    def test_custom_builder_receives_bytes_and_fmt(self, tmp_path: Path) -> None:
        """Builder is called once per embedded image with the raw
        bytes + normalized format + 1-based index."""
        orig = _doc_with_image(_png_data_uri(), alt="custom")
        out = tmp_path / "rt.docx"
        write_docx(orig, out)

        calls: list[tuple[int, str, int]] = []  # (len(data), fmt, index)

        def builder(data: bytes, fmt: str, index: int) -> str:
            calls.append((len(data), fmt, index))
            return f"artifact://img/{index}.{fmt}"

        reloaded = parse_docx(out, image_src_builder=builder)
        img = next(iter(find_by_type(reloaded, Image)), None)
        assert img is not None
        assert img.src == "artifact://img/1.png"
        # Exactly one invocation with the real PNG bytes + 1-based index.
        assert len(calls) == 1
        data_len, fmt, index = calls[0]
        assert fmt == "png"
        assert index == 1
        assert data_len == len(_minimal_png_bytes())

    def test_side_channel_byte_collection(self, tmp_path: Path) -> None:
        """The canonical use case: caller collects bytes in a dict
        keyed by the URI they return, so the AST stays lightweight
        while bytes land in a side store."""
        orig = _doc_with_image(_png_data_uri(), alt="side")
        out = tmp_path / "rt.docx"
        write_docx(orig, out)

        collected: dict[str, bytes] = {}

        def builder(data: bytes, fmt: str, index: int) -> str:
            uri = f"vfs://media/{index}.{fmt}"
            collected[uri] = data
            return uri

        reloaded = parse_docx(out, image_src_builder=builder)
        images = list(find_by_type(reloaded, Image))
        assert images
        assert set(collected) == {img.src for img in images}
        for img in images:
            assert collected[img.src] == _minimal_png_bytes()


class TestLibreOfficeShapeChecks:
    """Belt-and-suspenders structural assertions that catch common mistakes.

    These aren't a substitute for opening the file in LibreOffice, but
    they cover the failure modes that quietly pass in a round-trip test
    yet make Word show a red-X placeholder.
    """

    def test_r_namespace_declared_on_drawing_root(self) -> None:
        """If the ``r`` prefix isn't in scope where ``<a:blip r:embed=...>``
        is emitted, Word silently ignores the embed reference."""
        doc = _doc_with_image(_png_data_uri())
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml").decode("utf-8")
        assert f'xmlns:r="{R}"' in xml
        assert f'xmlns:wp="{WP}"' in xml
        assert f'xmlns:a="{A}"' in xml
        assert f'xmlns:pic="{PIC}"' in xml

    def test_w_drawing_element_is_child_of_w_r(self) -> None:
        doc = _doc_with_image(_png_data_uri())
        data = write_docx_bytes(doc)
        xml = _zip_read(data, "word/document.xml")
        root = etree.fromstring(xml)
        # drawing lives inside a w:r (a run) so inline text flow is preserved.
        drawings = root.findall(f".//{qn(W, 'r')}/{qn(W, 'drawing')}")
        assert drawings, "w:drawing must be a child of w:r"


class TestImageIsValidAgainstCalamineOrPython_docx:
    """Optional: if python-docx is importable, its reader should accept the file.

    python-docx is an independent OOXML implementation; if it opens and
    finds the picture, that's very strong evidence Word will too.
    """

    def test_python_docx_can_parse_our_output(self, tmp_path: Path) -> None:
        pytest.importorskip("docx")
        import docx  # ty: ignore[unresolved-import]

        doc = _doc_with_image(_png_data_uri(), alt="smoke", width=72.0, height=72.0)
        out = tmp_path / "smoke.docx"
        write_docx(doc, out)

        d = docx.Document(str(out))
        # python-docx surfaces inline images via the InlineShapes collection.
        assert len(d.inline_shapes) == 1, "python-docx should see exactly one inline shape"
