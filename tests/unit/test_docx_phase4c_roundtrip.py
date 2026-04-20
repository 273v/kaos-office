"""Tests for DOCX Phase 4C writer — multi-section emission + round-trip.

When ``doc.sections`` is populated, the writer emits one sectPr per
section: non-final sections land in an empty paragraph's ``<w:pPr>``
at the section boundary, the final section as the body-direct sectPr
(same shape Word produces).

These tests lock in:
- Backward compat (empty ``doc.sections`` still emits the Phase 4
  shape — one body-direct sectPr from ``metadata.page_setup``).
- Per-section geometry survives write → re-parse → re-write.
- The structural invariant: exactly ``len(sections)`` ``<w:sectPr>``
  elements are emitted, in document order, with matching break types.
- python-docx smoke test confirms an independent OOXML implementation
  accepts our multi-section output.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from kaos_content.model.blocks import Paragraph
from kaos_content.model.document import ContentDocument
from kaos_content.model.inlines import Text
from kaos_content.model.metadata import DocumentMetadata, PageSetup, Section
from lxml import etree  # ty: ignore[unresolved-import]

from kaos_office.docx.reader import parse_docx
from kaos_office.docx.writer import write_docx, write_docx_bytes
from kaos_office.ooxml.namespace import (
    W_BODY,
    W_PGSZ,
    W_PPR,
    W_SECTPR,
    W,
    qn,
)


def _body(data: bytes) -> etree._Element:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        xml = zf.read("word/document.xml")
    root = etree.fromstring(xml)
    body = root.find(W_BODY)
    assert body is not None
    return body


def _sectprs_in_order(body: etree._Element) -> list[etree._Element]:
    """Return every sectPr under body in document order — both pPr-nested
    and body-direct. This is what a conformant reader sees."""
    out: list[etree._Element] = []
    for el in body.iter():
        if el.tag == W_SECTPR:
            out.append(el)
    return out


def _para(text: str) -> Paragraph:
    return Paragraph(children=(Text(value=text),))


# --- Backward compat (empty sections) ------------------------------------


class TestBackwardCompatSingleSectPr:
    """No ``doc.sections`` => single body-direct sectPr (Phase 4 shape)."""

    def test_empty_sections_emits_single_sectpr(self) -> None:
        doc = ContentDocument(
            body=(_para("only paragraph"),),
            metadata=DocumentMetadata(
                page_setup=PageSetup(page_width_pt=612.0, page_height_pt=792.0)
            ),
        )
        assert doc.sections == ()
        body = _body(write_docx_bytes(doc))
        sectprs = _sectprs_in_order(body)
        assert len(sectprs) == 1
        # Sole sectPr is a body-direct child (not inside a pPr).
        assert sectprs[0].getparent() is body


# --- Multi-section emission ----------------------------------------------


class TestMultiSectionEmission:
    def test_two_sections_two_sectprs(self) -> None:
        """Exactly two sectPrs land in the package: one pPr-nested after
        block 2, one body-direct at body end."""
        doc = ContentDocument(
            body=(_para("p1"), _para("p2"), _para("p3"), _para("p4")),
            sections=(
                Section(
                    end_block_index=2,
                    page_setup=PageSetup(page_width_pt=612.0, page_height_pt=792.0),
                ),
                Section(
                    end_block_index=4,
                    page_setup=PageSetup(page_width_pt=792.0, page_height_pt=612.0),
                ),
            ),
        )
        body = _body(write_docx_bytes(doc))
        sectprs = _sectprs_in_order(body)
        assert len(sectprs) == 2
        # First is pPr-nested; second is body-direct.
        assert sectprs[0].getparent() is not None
        assert sectprs[0].getparent().tag == W_PPR
        assert sectprs[1].getparent() is body

        # Geometry distinct: first portrait, second landscape.
        def _dim(el: etree._Element, attr: str) -> int:
            pg = el.find(W_PGSZ)
            assert pg is not None
            return int(pg.get(qn(W, attr)) or "0")

        assert _dim(sectprs[0], "w") < _dim(sectprs[0], "h")  # portrait
        assert _dim(sectprs[1], "w") > _dim(sectprs[1], "h")  # landscape

    def test_break_type_emitted_when_non_default(self) -> None:
        """The writer omits <w:type> for the default ``nextPage`` and
        emits it verbatim for any other value."""
        doc = ContentDocument(
            body=(_para("p1"), _para("p2")),
            sections=(
                Section(end_block_index=1, break_type="continuous"),
                Section(end_block_index=2, break_type="nextPage"),
            ),
        )
        body = _body(write_docx_bytes(doc))
        sectprs = _sectprs_in_order(body)
        # Section 1 is continuous → <w:type w:val="continuous"/> present.
        t0 = sectprs[0].find(qn(W, "type"))
        assert t0 is not None
        assert t0.get(qn(W, "val")) == "continuous"
        # Section 2 is nextPage → <w:type/> omitted (OOXML default).
        assert sectprs[1].find(qn(W, "type")) is None

    def test_three_sections(self) -> None:
        doc = ContentDocument(
            body=(_para("a"), _para("b"), _para("c"), _para("d"), _para("e")),
            sections=(
                Section(end_block_index=2),
                Section(end_block_index=4),
                Section(end_block_index=5),
            ),
        )
        body = _body(write_docx_bytes(doc))
        assert len(_sectprs_in_order(body)) == 3


# --- Round-trip ----------------------------------------------------------


class TestRoundTrip:
    """write → parse_docx → write → parse_docx is idempotent for the
    section shape: section count + per-section geometry + break types."""

    def test_two_section_roundtrip_preserves_sections(self, tmp_path: Path) -> None:
        orig = ContentDocument(
            body=(_para("cover"), _para("body 1"), _para("body 2"), _para("tail")),
            sections=(
                Section(
                    end_block_index=1,
                    page_setup=PageSetup(page_width_pt=612.0, page_height_pt=792.0),
                    break_type="nextPage",
                ),
                Section(
                    end_block_index=4,
                    page_setup=PageSetup(page_width_pt=792.0, page_height_pt=612.0),
                    break_type="nextPage",
                ),
            ),
        )
        out = tmp_path / "two.docx"
        write_docx(orig, out)
        reloaded = parse_docx(out)
        assert len(reloaded.sections) == 2
        # end_block_index matches because the writer's extra empty
        # break paragraph has no runs, so the reader drops it and
        # the post-walk block count equals the pre-walk input count.
        assert reloaded.sections[0].end_block_index == 1
        assert reloaded.sections[1].end_block_index == 4
        # Geometry distinct: first portrait, second landscape.
        ps0 = reloaded.sections[0].page_setup
        ps1 = reloaded.sections[1].page_setup
        assert ps0 is not None and ps0.page_width_pt is not None
        assert ps1 is not None and ps1.page_width_pt is not None
        assert (
            ps0.page_width_pt is not None
            and ps0.page_height_pt is not None
            and ps0.page_width_pt < ps0.page_height_pt
        )
        assert (
            ps1.page_width_pt is not None
            and ps1.page_height_pt is not None
            and ps1.page_width_pt > ps1.page_height_pt
        )

    def test_roundtrip_idempotent(self, tmp_path: Path) -> None:
        """Second round trip produces the same section shape."""
        orig = ContentDocument(
            body=(_para("p1"), _para("p2"), _para("p3")),
            sections=(
                Section(
                    end_block_index=2,
                    page_setup=PageSetup(page_width_pt=612.0, page_height_pt=792.0),
                ),
                Section(
                    end_block_index=3,
                    page_setup=PageSetup(page_width_pt=792.0, page_height_pt=612.0),
                    break_type="continuous",
                ),
            ),
        )
        out1 = tmp_path / "r1.docx"
        write_docx(orig, out1)
        r1 = parse_docx(out1)
        out2 = tmp_path / "r2.docx"
        write_docx(r1, out2)
        r2 = parse_docx(out2)
        # Two round trips — structural shape stable.
        assert len(r2.sections) == len(r1.sections)
        for a, b in zip(r1.sections, r2.sections, strict=True):
            assert a.end_block_index == b.end_block_index
            assert a.break_type == b.break_type


# --- Headers / footers still attach to final sectPr ----------------------


class TestHeadersAttachToFinalSectPr:
    """Phase 4's header/footer refs continue to land on the body-direct
    (final) sectPr even in multi-section output."""

    def test_header_refs_on_final_sectpr(self) -> None:
        doc = ContentDocument(
            body=(_para("p1"), _para("p2")),
            headers={"default": (_para("H"),)},
            sections=(
                Section(end_block_index=1),
                Section(end_block_index=2),
            ),
        )
        body = _body(write_docx_bytes(doc))
        sectprs = _sectprs_in_order(body)
        final = sectprs[-1]
        # Final sectPr (body-direct) carries the header reference.
        refs = final.findall(qn(W, "headerReference"))
        assert len(refs) == 1
        # The non-final sectPr (pPr-nested) does not carry header refs —
        # those live on the final sectPr only in our current model.
        non_final = sectprs[0]
        assert non_final.findall(qn(W, "headerReference")) == []


# --- python-docx independent parser smoke test ---------------------------


class TestPythonDocxAcceptsMultiSection:
    """An independent OOXML implementation should open the file. If
    python-docx is absent, the test auto-skips."""

    def test_python_docx_parses_two_section_output(self, tmp_path: Path) -> None:
        pytest.importorskip("docx")
        import docx  # ty: ignore[unresolved-import]

        doc = ContentDocument(
            body=(_para("s1"), _para("s2")),
            sections=(
                Section(end_block_index=1),
                Section(end_block_index=2),
            ),
        )
        out = tmp_path / "twosec.docx"
        write_docx(doc, out)
        d = docx.Document(str(out))
        # python-docx surfaces sections via document.sections. Our
        # output should present exactly two sections (one per sectPr).
        assert len(d.sections) == 2
